
import io
import json
import math
import os
import re
import sqlite3
import hashlib
import secrets
from datetime import date, datetime, timedelta

import requests
import streamlit as st
from docx import Document
import xlsxwriter

# ============================================================
# THALASSEMIA SCREENING V5
# ============================================================
# 1) Hồ sơ bệnh nhân
# 2) Vòng 1: 20 câu hỏi
# 3) Chỉ nguy cơ CAO -> mở Vòng 2
# 4) Vòng 2:
#      - chọn tỉnh + phường/xã/đặc khu
#      - chọn khoảng độ cao
#      - CBC + đơn vị
#      - chuẩn hóa đơn vị
#      - Hb hiệu chỉnh theo độ cao (WHO 2024 prototype)
#      - Mentzer + phân tích sơ bộ
#      - lời khuyên
#      - gợi ý cơ sở y tế qua Google Places nếu có API key
#
# IMPORTANT:
# - Risk score hiện tại là prototype, chưa validation trên người Việt Nam.
# - "Khuyến nghị" là hỗ trợ sàng lọc, không chẩn đoán.
# - Q19/Q20 về tiếp cận xét nghiệm = 0 điểm.
# - Một số điện thoại = một hồ sơ; nhập lại sẽ ghi nhận lần cuối.
# - SQLite chỉ phù hợp prototype; Streamlit Cloud có thể reset filesystem.
# ============================================================


# ------------------------------------------------------------
# CONFIG
# ------------------------------------------------------------

st.set_page_config(
    page_title="Hệ thống Sàng lọc Thalassemia",
    page_icon="🩸",
    layout="wide",
)

DB_PATH = "thalassemia_patients.db"
CONSENT_VERSION = "THAL-RS-CONSENT-v1-2026-09-05"
ADMIN_DATA_URL = "https://raw.githubusercontent.com/open-admin-data/vietnam-administrative-divisions/main/data/hierarchy.json"
ADMIN_DATA_SOURCE_URL = "https://github.com/open-admin-data/vietnam-administrative-divisions"

ROUND1_MAX_SCORE = 20
ROUND1_HIGH_THRESHOLD = 8

FOLLOWUP_DAYS = 30

ALTITUDE_OPTIONS = {
    "<500 m": 0.0,
    "500–999 m": 0.4,
    "1.000–1.499 m": 0.8,
    "1.500–1.999 m": 1.1,
    "2.000–2.499 m": 1.4,
    "2.500–2.999 m": 1.8,
    "3.000–3.499 m": 2.1,
    "3.500–3.999 m": 2.5,
    "4.000–4.499 m": 2.9,
    "4.500–4.999 m": 3.3,
}

ALTITUDE_LOOKUP_URL = "https://elevationfinder.net/"


# ------------------------------------------------------------
# GOOGLE KEY
# ------------------------------------------------------------

def get_google_key():
    try:
        value = st.secrets["google"]["maps_api_key"]
        if value:
            return str(value).strip()
    except Exception:
        pass

    try:
        value = st.secrets["GOOGLE_MAPS_API_KEY"]
        if value:
            return str(value).strip()
    except Exception:
        pass

    return os.environ.get(
        "GOOGLE_MAPS_API_KEY",
        "",
    ).strip()


GOOGLE_API_KEY = get_google_key()


# ------------------------------------------------------------
# DATABASE
# ------------------------------------------------------------

PBKDF2_ITERATIONS = 310_000

# Tài khoản quản trị mặc định theo yêu cầu của chủ hệ thống.
DEFAULT_ADMIN_USERNAME = "daclau239"
DEFAULT_ADMIN_PASSWORD = "23092002"
DEFAULT_ADMIN_FULL_NAME = "Quản trị viên hệ thống"
DEFAULT_ADMIN_EMAIL = "admin@thalassemia.local"


def get_db():
    conn = sqlite3.connect(
        DB_PATH,
        check_same_thread=False,
    )

    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS patient_profiles (
            phone TEXT PRIMARY KEY,
            full_name TEXT NOT NULL,
            birth_date TEXT NOT NULL,
            gender TEXT NOT NULL,
            current_address TEXT NOT NULL,
            province TEXT NOT NULL,
            commune TEXT NOT NULL,
            research_consent INTEGER NOT NULL DEFAULT 0,
            consent_version TEXT,
            consent_at TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        )
        """
    )

    # Tài khoản quản trị / người được phê duyệt.
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS user_accounts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL UNIQUE,
            full_name TEXT NOT NULL,
            email TEXT NOT NULL UNIQUE,
            password_hash TEXT NOT NULL,
            password_salt TEXT NOT NULL,
            role TEXT NOT NULL CHECK(role IN ('admin', 'staff')),
            status TEXT NOT NULL CHECK(status IN ('pending', 'approved', 'rejected', 'disabled')),
            created_at TEXT NOT NULL,
            approved_by TEXT,
            approved_at TEXT,
            last_login_at TEXT
        )
        """
    )

    # Nhật ký từng lần sàng lọc: giữ lịch sử theo lượt, tách khỏi hồ sơ hiện tại.
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS screening_records (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            phone TEXT NOT NULL,
            screening_at TEXT NOT NULL,
            entered_by_username TEXT,
            entry_mode TEXT NOT NULL CHECK(entry_mode IN ('self', 'assisted')),
            consent_version TEXT NOT NULL,
            consent_at TEXT NOT NULL,
            round1_score INTEGER NOT NULL,
            round1_category TEXT NOT NULL,
            round1_conclusion TEXT,
            round1_reasons TEXT,
            answers_json TEXT,
            round2_completed INTEGER NOT NULL DEFAULT 0,
            altitude_choice TEXT,
            altitude_adjustment REAL,
            hb REAL,
            hb_adjusted REAL,
            mcv REAL,
            mch REAL,
            rbc REAL,
            rdw REAL,
            mentzer REAL,
            round2_score INTEGER,
            round2_category TEXT,
            round2_conclusion TEXT,
            round2_reasons TEXT,
            findings_json TEXT,
            advice_json TEXT,
            FOREIGN KEY(phone) REFERENCES patient_profiles(phone)
        )
        """
    )

    # Migrate prototype databases created before consent fields existed.
    existing = {
        row[1]
        for row in conn.execute(
            "PRAGMA table_info(patient_profiles)"
        ).fetchall()
    }
    migrations = [
        ("research_consent", "INTEGER NOT NULL DEFAULT 0"),
        ("consent_version", "TEXT"),
        ("consent_at", "TEXT"),
    ]
    for column, definition in migrations:
        if column not in existing:
            conn.execute(
                f"ALTER TABLE patient_profiles ADD COLUMN {column} {definition}"
            )

    conn.commit()
    return conn


def get_secret_value(*paths, env_name=""):
    for path in paths:
        try:
            value = st.secrets
            for part in path.split("."):
                value = value[part]
            if value:
                return str(value).strip()
        except Exception:
            pass

    if env_name:
        return os.environ.get(env_name, "").strip()
    return ""


def hash_password(password, salt_hex=None):
    if salt_hex:
        salt = bytes.fromhex(salt_hex)
        salt_value = salt_hex
    else:
        salt = secrets.token_bytes(16)
        salt_value = salt.hex()

    digest = hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        salt,
        PBKDF2_ITERATIONS,
    ).hex()

    return digest, salt_value


def verify_password(password, stored_hash, stored_salt):
    digest, _ = hash_password(password, stored_salt)
    return secrets.compare_digest(digest, stored_hash)


def valid_email(email):
    return bool(
        re.fullmatch(
            r"[^@\s]+@[^@\s]+\.[^@\s]+",
            (email or "").strip(),
        )
    )


def valid_username(username):
    return bool(
        re.fullmatch(
            r"[A-Za-z0-9_.-]{4,32}",
            (username or "").strip(),
        )
    )


def ensure_default_admin():
    """Tạo admin mặc định một lần nếu database chưa có tài khoản admin."""
    conn = get_db()
    row = conn.execute(
        "SELECT id FROM user_accounts WHERE username = ? LIMIT 1",
        (DEFAULT_ADMIN_USERNAME,),
    ).fetchone()
    if row:
        conn.close()
        return

    admin_row = conn.execute(
        "SELECT id FROM user_accounts WHERE role = 'admin' LIMIT 1"
    ).fetchone()
    if admin_row:
        conn.close()
        return

    now = datetime.now().isoformat(timespec="seconds")
    password_hash, password_salt = hash_password(DEFAULT_ADMIN_PASSWORD)
    conn.execute(
        """
        INSERT INTO user_accounts
        (username, full_name, email, password_hash, password_salt,
         role, status, created_at, approved_by, approved_at)
        VALUES (?, ?, ?, ?, ?, 'admin', 'approved', ?, ?, ?)
        """,
        (
            DEFAULT_ADMIN_USERNAME,
            DEFAULT_ADMIN_FULL_NAME,
            DEFAULT_ADMIN_EMAIL,
            password_hash,
            password_salt,
            now,
            DEFAULT_ADMIN_USERNAME,
            now,
        ),
    )
    conn.commit()
    conn.close()


def admin_exists():
    conn = get_db()
    row = conn.execute(
        "SELECT 1 FROM user_accounts WHERE role = 'admin' LIMIT 1"
    ).fetchone()
    conn.close()
    return row is not None


def create_admin_account(username, full_name, email, password):
    username = username.strip()
    email = email.strip().lower()
    now = datetime.now().isoformat(timespec="seconds")
    password_hash, password_salt = hash_password(password)

    conn = get_db()
    try:
        conn.execute(
            """
            INSERT INTO user_accounts
            (username, full_name, email, password_hash, password_salt,
             role, status, created_at, approved_by, approved_at)
            VALUES (?, ?, ?, ?, ?, 'admin', 'approved', ?, ?, ?)
            """,
            (
                username,
                full_name.strip(),
                email,
                password_hash,
                password_salt,
                now,
                username,
                now,
            ),
        )
        conn.commit()
        return True, ""
    except sqlite3.IntegrityError as exc:
        conn.rollback()
        return False, f"Tài khoản/email đã tồn tại hoặc dữ liệu không hợp lệ: {exc}"
    finally:
        conn.close()


def register_staff_account(username, full_name, email, password):
    username = username.strip()
    email = email.strip().lower()
    now = datetime.now().isoformat(timespec="seconds")
    password_hash, password_salt = hash_password(password)

    conn = get_db()
    try:
        conn.execute(
            """
            INSERT INTO user_accounts
            (username, full_name, email, password_hash, password_salt,
             role, status, created_at)
            VALUES (?, ?, ?, ?, ?, 'staff', 'pending', ?)
            """,
            (
                username,
                full_name.strip(),
                email,
                password_hash,
                password_salt,
                now,
            ),
        )
        conn.commit()
        return True, ""
    except sqlite3.IntegrityError:
        conn.rollback()
        return False, "Tên đăng nhập hoặc email đã được sử dụng."
    finally:
        conn.close()


def authenticate_user(login_value, password):
    login_value = (login_value or "").strip().lower()
    conn = get_db()
    row = conn.execute(
        """
        SELECT id, username, full_name, email, password_hash,
               password_salt, role, status
        FROM user_accounts
        WHERE lower(username) = ? OR lower(email) = ?
        LIMIT 1
        """,
        (login_value, login_value),
    ).fetchone()

    if not row:
        conn.close()
        return None, "Sai tên đăng nhập/email hoặc mật khẩu."

    if not verify_password(password, row[4], row[5]):
        conn.close()
        return None, "Sai tên đăng nhập/email hoặc mật khẩu."

    if row[7] != "approved":
        conn.close()
        if row[7] == "pending":
            return None, "Tài khoản đang chờ quản trị viên phê duyệt."
        if row[7] == "disabled":
            return None, "Tài khoản đã bị vô hiệu hóa."
        return None, "Tài khoản chưa được phép truy cập."

    now = datetime.now().isoformat(timespec="seconds")
    conn.execute(
        "UPDATE user_accounts SET last_login_at = ? WHERE id = ?",
        (now, row[0]),
    )
    conn.commit()
    conn.close()

    return {
        "id": row[0],
        "username": row[1],
        "full_name": row[2],
        "email": row[3],
        "role": row[6],
        "status": row[7],
    }, ""


def list_user_accounts():
    conn = get_db()
    rows = conn.execute(
        """
        SELECT id, username, full_name, email, role, status,
               created_at, approved_by, approved_at, last_login_at
        FROM user_accounts
        ORDER BY
            CASE role WHEN 'admin' THEN 0 ELSE 1 END,
            CASE status WHEN 'pending' THEN 0 ELSE 1 END,
            created_at DESC
        """
    ).fetchall()
    conn.close()
    return rows


def update_staff_status(user_id, status, approved_by):
    if status not in {"approved", "rejected", "disabled", "pending"}:
        raise ValueError("Trạng thái không hợp lệ.")

    conn = get_db()
    now = datetime.now().isoformat(timespec="seconds")
    if status == "approved":
        conn.execute(
            """
            UPDATE user_accounts
            SET status = ?, approved_by = ?, approved_at = ?
            WHERE id = ? AND role = 'staff'
            """,
            (status, approved_by, now, user_id),
        )
    else:
        conn.execute(
            """
            UPDATE user_accounts
            SET status = ?
            WHERE id = ? AND role = 'staff'
            """,
            (status, user_id),
        )
    conn.commit()
    conn.close()


def list_patient_profiles_for_staff():
    conn = get_db()
    rows = conn.execute(
        """
        SELECT phone, full_name, birth_date, gender,
               current_address, province, commune,
               research_consent, consent_version, consent_at,
               created_at, updated_at
        FROM patient_profiles
        WHERE research_consent = 1
        ORDER BY updated_at DESC
        """
    ).fetchall()
    conn.close()
    return rows


def phone_exists(phone):
    conn = get_db()

    row = conn.execute(
        """
        SELECT 1
        FROM patient_profiles
        WHERE phone = ?
        """,
        (phone,),
    ).fetchone()

    conn.close()
    return row is not None


def upsert_patient(profile):
    conn = get_db()

    now = datetime.now().isoformat(
        timespec="seconds"
    )

    exists = conn.execute(
        """
        SELECT 1
        FROM patient_profiles
        WHERE phone = ?
        """,
        (profile["phone"],),
    ).fetchone()

    if exists:
        conn.execute(
            """
            UPDATE patient_profiles
            SET
                full_name = ?,
                birth_date = ?,
                gender = ?,
                current_address = ?,
                province = ?,
                commune = ?,
                research_consent = ?,
                consent_version = ?,
                consent_at = ?,
                updated_at = ?
            WHERE phone = ?
            """,
            (
                profile["full_name"],
                profile["birth_date"],
                profile["gender"],
                profile["current_address"],
                profile["province"],
                profile["commune"],
                1,
                CONSENT_VERSION,
                profile["consent_at"],
                now,
                profile["phone"],
            ),
        )
        action = "updated"
    else:
        conn.execute(
            """
            INSERT INTO patient_profiles (
                phone,
                full_name,
                birth_date,
                gender,
                current_address,
                province,
                commune,
                research_consent,
                consent_version,
                consent_at,
                created_at,
                updated_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                profile["phone"],
                profile["full_name"],
                profile["birth_date"],
                profile["gender"],
                profile["current_address"],
                profile["province"],
                profile["commune"],
                1,
                CONSENT_VERSION,
                profile["consent_at"],
                now,
                now,
            ),
        )
        action = "inserted"

    conn.commit()
    conn.close()

    return action


def create_screening_record(
    patient,
    answers,
    score1,
    category1,
    conclusion1,
    reasons1,
    entry_mode,
    entered_by_username=None,
):
    """Lưu một lượt sàng lọc sau khi hoàn thành Vòng 1."""
    now = datetime.now().isoformat(timespec="seconds")
    conn = get_db()
    cur = conn.execute(
        """
        INSERT INTO screening_records (
            phone, screening_at, entered_by_username, entry_mode,
            consent_version, consent_at, round1_score, round1_category,
            round1_conclusion, round1_reasons, answers_json
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            patient["phone"],
            now,
            entered_by_username,
            entry_mode,
            CONSENT_VERSION,
            patient["consent_at"],
            int(score1),
            category1,
            conclusion1,
            json.dumps(reasons1, ensure_ascii=False),
            json.dumps(answers, ensure_ascii=False),
        ),
    )
    record_id = cur.lastrowid
    conn.commit()
    conn.close()
    return record_id


def update_screening_round2(record_id, r2):
    if not record_id:
        return

    conn = get_db()
    conn.execute(
        """
        UPDATE screening_records
        SET round2_completed = 1,
            altitude_choice = ?,
            altitude_adjustment = ?,
            hb = ?,
            hb_adjusted = ?,
            mcv = ?,
            mch = ?,
            rbc = ?,
            rdw = ?,
            mentzer = ?,
            round2_score = ?,
            round2_category = ?,
            round2_conclusion = ?,
            round2_reasons = ?,
            findings_json = ?,
            advice_json = ?
        WHERE id = ?
        """,
        (
            r2["altitude_choice"],
            r2["adjustment"],
            r2["hb"],
            r2["hb_adjusted"],
            r2["mcv"],
            r2["mch"],
            r2["rbc"],
            r2["rdw"],
            r2["mentzer"],
            r2["score"],
            r2["category"],
            r2["conclusion"],
            json.dumps(r2["reasons"], ensure_ascii=False),
            json.dumps(r2["findings"], ensure_ascii=False),
            json.dumps(r2["advice"], ensure_ascii=False),
            int(record_id),
        ),
    )
    conn.commit()
    conn.close()


def list_screening_records_for_staff():
    conn = get_db()
    rows = conn.execute(
        """
        SELECT
            s.id, s.screening_at, s.entered_by_username, s.entry_mode,
            p.phone, p.full_name, p.birth_date, p.gender,
            p.current_address, p.province, p.commune,
            s.consent_version, s.consent_at,
            s.round1_score, s.round1_category,
            s.round2_completed, s.altitude_choice, s.altitude_adjustment,
            s.hb, s.hb_adjusted, s.mcv, s.mch, s.rbc, s.rdw, s.mentzer,
            s.round2_score, s.round2_category, s.round2_conclusion,
            s.round1_conclusion
        FROM screening_records s
        JOIN patient_profiles p ON p.phone = s.phone
        WHERE p.research_consent = 1
        ORDER BY s.screening_at DESC
        """
    ).fetchall()
    conn.close()
    return rows


def export_screening_xlsx(patient_rows, screening_rows):
    """Tạo một file Excel nhiều sheet, không lưu file tạm trên server."""
    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output, {"in_memory": True})

    header_fmt = workbook.add_format({
        "bold": True,
        "bg_color": "#D9EAF7",
        "border": 1,
        "align": "center",
        "valign": "vcenter",
        "text_wrap": True,
    })
    cell_fmt = workbook.add_format({"border": 1, "valign": "top"})
    date_fmt = workbook.add_format({"border": 1, "num_format": "dd/mm/yyyy hh:mm"})

    # Sheet 1: hồ sơ hiện tại
    ws = workbook.add_worksheet("Ho_so_nguoi_tham_gia")
    patient_headers = [
        "Số điện thoại", "Họ và tên", "Ngày sinh", "Giới tính",
        "Địa chỉ hiện tại", "Tỉnh/thành", "Phường/xã/đặc khu",
        "Đồng ý nghiên cứu", "Phiên bản consent", "Thời điểm đồng ý",
        "Tạo lúc", "Cập nhật lúc",
    ]
    for c, h in enumerate(patient_headers):
        ws.write(0, c, h, header_fmt)
    for r, row in enumerate(patient_rows, start=1):
        values = [
            row[0], row[1], row[2], row[3], row[4], row[5], row[6],
            "Có" if row[7] else "Không", row[8] or "", row[9] or "", row[10], row[11],
        ]
        for c, value in enumerate(values):
            fmt = date_fmt if c in (9, 10, 11) and isinstance(value, datetime) else cell_fmt
            ws.write(r, c, value, fmt)

    ws.freeze_panes(1, 0)
    widths = [15, 24, 13, 12, 32, 20, 26, 16, 25, 21, 21, 21]
    for c, w in enumerate(widths):
        ws.set_column(c, c, w)
    ws.autofilter(0, 0, max(len(patient_rows), 1), len(patient_headers)-1)

    # Sheet 2: lịch sử sàng lọc
    ws2 = workbook.add_worksheet("Lich_su_sang_loc")
    screening_headers = [
        "ID lượt sàng lọc", "Thời điểm", "Người nhập", "Hình thức nhập",
        "Số điện thoại", "Họ và tên", "Ngày sinh", "Giới tính",
        "Địa chỉ", "Tỉnh/thành", "Phường/xã/đặc khu", "Phiên bản consent",
        "Thời điểm đồng ý", "Điểm Vòng 1", "Nguy cơ Vòng 1",
        "Đã hoàn thành Vòng 2", "Khoảng độ cao", "Hiệu chỉnh Hb (g/dL)",
        "Hb (g/dL)", "Hb sau hiệu chỉnh (g/dL)", "MCV (fL)", "MCH (pg)",
        "RBC (T/L)", "RDW-CV (%)", "Mentzer Index", "Điểm CBC",
        "Nguy cơ Vòng 2", "Kết luận Vòng 2", "Kết luận Vòng 1",
    ]
    for c, h in enumerate(screening_headers):
        ws2.write(0, c, h, header_fmt)

    for r, row in enumerate(screening_rows, start=1):
        values = [
            row[0], row[1], row[2] or "",
            "Tự nhập" if row[3] == "self" else "Nhập giúp người tham gia",
            row[4], row[5], row[6], row[7], row[8], row[9], row[10],
            row[11], row[12], row[13], row[14],
            "Có" if row[15] else "Chưa", row[16] or "", row[17] if row[17] is not None else "",
            row[18] if row[18] is not None else "", row[19] if row[19] is not None else "",
            row[20] if row[20] is not None else "", row[21] if row[21] is not None else "",
            row[22] if row[22] is not None else "", row[23] if row[23] is not None else "",
            row[24] if row[24] is not None else "", row[25] if row[25] is not None else "",
            row[26] or "", row[27] or "", row[28] or "",
        ]
        for c, value in enumerate(values):
            fmt = date_fmt if c in (1, 12) and isinstance(value, datetime) else cell_fmt
            ws2.write(r, c, value, fmt)

    ws2.freeze_panes(1, 0)
    for c in range(len(screening_headers)):
        ws2.set_column(c, c, 20)
    ws2.set_column(4, 4, 15)
    ws2.set_column(5, 5, 24)
    ws2.set_column(8, 8, 32)
    ws2.set_column(9, 10, 22)
    ws2.set_column(27, 28, 38)
    ws2.autofilter(0, 0, max(len(screening_rows), 1), len(screening_headers)-1)

    workbook.close()
    output.seek(0)
    return output


# ------------------------------------------------------------
# ACCESS CONTROL / ADMIN CONSOLE
# ------------------------------------------------------------

ensure_default_admin()


def current_auth_user():
    return st.session_state.get("auth_user")


def logout_user():
    st.session_state.pop("auth_user", None)


def render_auth_sidebar():
    user = current_auth_user()

    st.sidebar.header("🔐 TÀI KHOẢN & QUYỀN TRUY CẬP")

    if user:
        role_label = "Quản trị viên" if user["role"] == "admin" else "Nhân sự được duyệt"
        st.sidebar.success(
            f"Đang đăng nhập: **{user['full_name']}**\n\n{role_label}"
        )

        if user["role"] == "admin":
            page = st.sidebar.radio(
                "Khu vực làm việc",
                ["📝 Nhập sàng lọc", "🛡️ Quản trị hệ thống"],
                key="auth_page_admin",
            )
        else:
            page = "📝 Nhập sàng lọc"
            st.sidebar.info(
                "Bạn có thể nhập hồ sơ/sàng lọc giúp người tham gia. "
                "Danh sách hồ sơ chỉ xem được ở khu vực quản trị có quyền."
            )

        st.session_state["auth_page"] = page

        if st.sidebar.button("🚪 Đăng xuất", use_container_width=True):
            logout_user()
            st.session_state.pop("auth_page", None)
            st.rerun()
        return

    st.sidebar.caption(
        "Người tham gia có thể tự nhập hồ sơ.\n"
        "Quản trị viên/nhân sự được duyệt đăng nhập để quản lý hoặc nhập giúp người tham gia."
    )

    auth_mode = st.sidebar.radio(
        "",
        ["Tài khoản quản trị", "Đăng ký nhân sự"],
        key="auth_mode",
    )

    if auth_mode == "Tài khoản quản trị":
        with st.sidebar.form("login_form"):
            login_value = st.text_input("Tên đăng nhập hoặc email")
            password = st.text_input("Mật khẩu", type="password")
            login_submit = st.form_submit_button(
                "🔑 ĐĂNG NHẬP",
                use_container_width=True,
            )

        if login_submit:
            user_result, message = authenticate_user(login_value, password)
            if user_result:
                st.session_state["auth_user"] = user_result
                st.session_state["auth_page"] = "📝 Nhập sàng lọc"
                st.rerun()
            else:
                st.sidebar.error(message)

        st.sidebar.info(
            "Tài khoản quản trị đã được cấu hình sẵn cho hệ thống. "
            "Đăng nhập bằng tài khoản được cấp cho quản trị viên."
        )
    else:
        with st.sidebar.form("staff_register_form"):
            staff_username = st.text_input("Tên đăng nhập")
            staff_full_name = st.text_input("Họ và tên")
            staff_email = st.text_input("Email")
            staff_password = st.text_input("Mật khẩu", type="password")
            staff_password2 = st.text_input("Nhập lại mật khẩu", type="password")
            register_submit = st.form_submit_button(
                "📝 GỬI YÊU CẦU TẠO TÀI KHOẢN",
                use_container_width=True,
            )

        if register_submit:
            if not valid_username(staff_username):
                st.sidebar.error("Tên đăng nhập 4–32 ký tự, không có khoảng trắng.")
            elif not staff_full_name.strip():
                st.sidebar.error("Vui lòng nhập họ và tên.")
            elif not valid_email(staff_email):
                st.sidebar.error("Email chưa đúng định dạng.")
            elif len(staff_password) < 8:
                st.sidebar.error("Mật khẩu phải có ít nhất 8 ký tự.")
            elif staff_password != staff_password2:
                st.sidebar.error("Hai mật khẩu không khớp.")
            else:
                ok, msg = register_staff_account(
                    staff_username, staff_full_name, staff_email, staff_password
                )
                if ok:
                    st.sidebar.success(
                        "✅ Đã gửi tài khoản. Quản trị viên phải phê duyệt trước khi đăng nhập."
                    )
                else:
                    st.sidebar.error(msg)

        st.sidebar.caption(
            "Tài khoản nhân sự không được xem dữ liệu người tham gia cho đến khi quản trị viên phê duyệt."
        )


def render_admin_console(user):
    st.header("🛡️ QUẢN TRỊ HỆ THỐNG")
    st.success(
        f"Xin chào **{user['full_name']}** — quyền: "
        f"{'Quản trị viên' if user['role'] == 'admin' else 'Nhân sự được duyệt'}"
    )

    patients = list_patient_profiles_for_staff()
    users = list_user_accounts()
    screening_rows = list_screening_records_for_staff()

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.metric("Hồ sơ đã đồng ý", len(patients))
    with c2:
        pending_count = sum(1 for row in users if row[5] == "pending" and row[4] == "staff")
        st.metric("Chờ phê duyệt", pending_count)
    with c3:
        staff_count = sum(1 for row in users if row[4] == "staff" and row[5] == "approved")
        st.metric("Nhân sự được duyệt", staff_count)
    with c4:
        st.metric("Lượt sàng lọc đã lưu", len(screening_rows))

    if user["role"] == "admin":
        st.subheader("👥 Phê duyệt tài khoản nhân sự")
        pending_users = [row for row in users if row[4] == "staff" and row[5] == "pending"]
        if not pending_users:
            st.info("Hiện không có tài khoản nào đang chờ phê duyệt.")
        else:
            for row in pending_users:
                with st.container(border=True):
                    a, b, c = st.columns([2, 2, 1])
                    with a:
                        st.write(f"**{row[2]}**")
                        st.caption(f"Username: {row[1]} · Email: {row[3]}")
                    with b:
                        st.caption(f"Đăng ký: {row[6]}")
                    with c:
                        b1, b2 = st.columns(2)
                        with b1:
                            if st.button("✅ Duyệt", key=f"approve_{row[0]}"):
                                update_staff_status(row[0], "approved", user["username"])
                                st.rerun()
                        with b2:
                            if st.button("❌ Từ chối", key=f"reject_{row[0]}"):
                                update_staff_status(row[0], "rejected", user["username"])
                                st.rerun()

        st.subheader("⚙️ Tài khoản nhân sự")
        staff_rows = [row for row in users if row[4] == "staff"]
        for row in staff_rows:
            with st.container(border=True):
                s1, s2, s3 = st.columns([2, 2, 1])
                with s1:
                    st.write(f"**{row[2]}**")
                    st.caption(f"{row[1]} · {row[3]}")
                with s2:
                    st.write(f"Trạng thái: **{row[5]}**")
                with s3:
                    if row[5] == "approved":
                        if st.button("Vô hiệu hóa", key=f"disable_{row[0]}"):
                            update_staff_status(row[0], "disabled", user["username"])
                            st.rerun()
                    elif row[5] in {"disabled", "rejected"}:
                        if st.button("Mở lại", key=f"enable_{row[0]}"):
                            update_staff_status(row[0], "approved", user["username"])
                            st.rerun()

    st.divider()
    st.subheader("📊 DỮ LIỆU NGƯỜI THAM GIA — DẠNG BẢNG")
    st.caption(
        "🔒 Chỉ quản trị viên và nhân sự đã được quản trị viên phê duyệt mới xem được dữ liệu này. "
        "Dữ liệu gồm hồ sơ hiện tại và lịch sử từng lượt sàng lọc đã đồng ý."
    )

    tab1, tab2 = st.tabs(["👤 Hồ sơ hiện tại", "🧪 Lịch sử sàng lọc"])

    with tab1:
        patient_columns = [
            "Số điện thoại", "Họ tên", "Ngày sinh", "Giới tính",
            "Địa chỉ hiện tại", "Tỉnh/thành", "Phường/xã/đặc khu",
            "Đồng ý nghiên cứu", "Phiên bản consent", "Thời điểm đồng ý",
            "Tạo lúc", "Cập nhật lúc",
        ]
        patient_table = []
        for row in patients:
            patient_table.append({
                "Số điện thoại": row[0],
                "Họ tên": row[1],
                "Ngày sinh": row[2],
                "Giới tính": row[3],
                "Địa chỉ hiện tại": row[4],
                "Tỉnh/thành": row[5],
                "Phường/xã/đặc khu": row[6],
                "Đồng ý nghiên cứu": "Có" if row[7] else "Không",
                "Phiên bản consent": row[8] or "",
                "Thời điểm đồng ý": row[9] or "",
                "Tạo lúc": row[10],
                "Cập nhật lúc": row[11],
            })
        if patient_table:
            st.dataframe(patient_table, use_container_width=True, hide_index=True)
        else:
            st.info("Chưa có hồ sơ nào đã đồng ý tham gia.")

    with tab2:
        screening_columns = [
            "ID", "Thời điểm", "Người nhập", "Hình thức nhập", "Số điện thoại",
            "Họ tên", "Tỉnh/thành", "Phường/xã/đặc khu", "Điểm Vòng 1",
            "Nguy cơ Vòng 1", "Vòng 2", "Độ cao", "Hb", "Hb sau hiệu chỉnh",
            "MCV", "MCH", "RBC", "RDW", "Mentzer", "Điểm CBC",
            "Nguy cơ Vòng 2", "Kết luận",
        ]
        screening_table = []
        for row in screening_rows:
            screening_table.append({
                "ID": row[0],
                "Thời điểm": row[1],
                "Người nhập": row[2] or "",
                "Hình thức nhập": "Tự nhập" if row[3] == "self" else "Nhập giúp người tham gia",
                "Số điện thoại": row[4],
                "Họ tên": row[5],
                "Tỉnh/thành": row[9],
                "Phường/xã/đặc khu": row[10],
                "Điểm Vòng 1": row[13],
                "Nguy cơ Vòng 1": row[14],
                "Vòng 2": "Có" if row[15] else "Chưa",
                "Độ cao": row[16] or "",
                "Hb": row[18] if row[18] is not None else "",
                "Hb sau hiệu chỉnh": row[19] if row[19] is not None else "",
                "MCV": row[20] if row[20] is not None else "",
                "MCH": row[21] if row[21] is not None else "",
                "RBC": row[22] if row[22] is not None else "",
                "RDW": row[23] if row[23] is not None else "",
                "Mentzer": row[24] if row[24] is not None else "",
                "Điểm CBC": row[25] if row[25] is not None else "",
                "Nguy cơ Vòng 2": row[26] or "",
                "Kết luận": row[27] or row[28] or "",
            })
        if screening_table:
            st.dataframe(screening_table, use_container_width=True, hide_index=True)
        else:
            st.info("Chưa có lượt sàng lọc nào được lưu.")

    if patients or screening_rows:
        excel_data = export_screening_xlsx(patients, screening_rows)
        st.download_button(
            "📊 XUẤT DỮ LIỆU EXCEL (.xlsx)",
            data=excel_data.getvalue(),
            file_name=f"Thalassemia_du_lieu_{date.today().isoformat()}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )
        st.caption(
            "File Excel gồm 2 sheet: Hồ sơ hiện tại và Lịch sử sàng lọc. "
            "Không xuất mật khẩu/tài khoản nhân sự."
        )


# ------------------------------------------------------------
# CURRENT VIETNAM ADMINISTRATIVE DATA (34 PROVINCES / 3,321 COMMUNES)
# ------------------------------------------------------------

@st.cache_data(ttl=86400, show_spinner=False)
def load_admin_hierarchy():
    """Load the current 2-level administrative hierarchy.

    Source dataset follows Vietnam's post-2025 structure: 34 provincial
    units directly administering 3,321 commune/ward/special-area units.
    The prototype deliberately does NOT fall back to manual commune entry.
    """
    response = requests.get(ADMIN_DATA_URL, timeout=20)
    response.raise_for_status()
    payload = response.json()

    records = payload.get("data", payload)
    if not isinstance(records, list) or not records:
        raise ValueError("Dữ liệu địa giới không đúng định dạng.")

    provinces = {}
    for province in records:
        province_name = province.get("name", {}).get("local")
        if not province_name:
            continue

        wards = []
        for ward in province.get("ward", []) or province.get("wards", []):
            ward_name = ward.get("name", {}).get("local")
            if ward_name:
                wards.append(ward_name)

        provinces[province_name] = sorted(set(wards), key=str.casefold)

    if len(provinces) != 34:
        raise ValueError(
            f"Dữ liệu địa giới hiện trả về {len(provinces)} tỉnh/thành, không phải 34."
        )

    total_communes = sum(len(items) for items in provinces.values())
    if total_communes != 3321:
        raise ValueError(
            f"Dữ liệu địa giới hiện có {total_communes} đơn vị cấp xã, không phải 3.321."
        )

    return provinces


# ------------------------------------------------------------
# GENERAL HELPERS
# ------------------------------------------------------------

def calculate_age(birth_date):
    today = date.today()

    return (
        today.year
        - birth_date.year
        - (
            (today.month, today.day)
            < (birth_date.month, birth_date.day)
        )
    )


def safe_filename(text):
    text = re.sub(
        r"[^0-9A-Za-zÀ-ỹĐđ _-]",
        "_",
        (text or "").strip(),
    )
    return text.strip(" _") or "nguoi_sang_loc"


def reset_results():
    st.session_state.pop("screening_id", None)
    for key in list(st.session_state.keys()):
        if (
            key.startswith("round1_")
            or key.startswith("round2_")
            or key.startswith("low_cbc_")
            or key.startswith("google_")
        ):
            del st.session_state[key]


# ------------------------------------------------------------
# ALTITUDE
# ------------------------------------------------------------

def altitude_adjustment_from_choice(choice):
    return ALTITUDE_OPTIONS.get(
        choice,
        0.0,
    )


def altitude_selector(key_prefix):
    choice = st.radio(
        "Chọn khoảng độ cao nơi đang sinh sống",
        list(ALTITUDE_OPTIONS.keys()),
        index=0,
        key=f"{key_prefix}_choice",
    )

    st.link_button(
        "🔎 Tra cứu độ cao nơi ở",
        ALTITUDE_LOOKUP_URL,
    )

    adjustment = altitude_adjustment_from_choice(
        choice
    )

    st.info(
        f"Hiệu chỉnh Hb tham khảo: **-{adjustment:.1f} g/dL**"
    )

    if choice.startswith(
        (
            "2.500",
            "3.000",
            "3.500",
            "4.000",
            "4.500",
        )
    ):
        st.warning(
            "Ở độ cao ≥2.500 m, WHO lưu ý mức độ không chắc chắn "
            "của hiệu chỉnh cao hơn."
        )

    return choice, adjustment


# ------------------------------------------------------------
# CBC UNIT CONVERSION
# ------------------------------------------------------------

def hb_to_g_dl(value, unit):
    if unit == "g/dL":
        return float(value)

    if unit == "g/L":
        return float(value) / 10

    raise ValueError("Đơn vị Hb không hợp lệ.")


def rbc_to_t_l(value, unit):
    # Numerically equivalent:
    # T/L = 10^12/L = 10^6/µL
    if unit in (
        "T/L",
        "10^12/L",
        "10^6/µL",
    ):
        return float(value)

    raise ValueError("Đơn vị RBC không hợp lệ.")


# ------------------------------------------------------------
# ROUND 1 SCORE
# ------------------------------------------------------------

def calculate_round1_score(a):
    score = 0
    reasons = []

    weighted_items = [
        ("q1", 3, "Có người thân/dòng họ mắc Thalassemia"),
        ("q2", 3, "Có người thân/dòng họ mang gen Thalassemia/hemoglobinopathy"),
        ("q3", 1, "Cha/mẹ từng xét nghiệm Thalassemia/hemoglobinopathy"),
        ("q4", 2, "Anh/chị/em từng thiếu máu hoặc hồng cầu nhỏ"),
        ("q5", 2, "Gia đình có trẻ từng truyền máu nhiều lần/định kỳ"),
        ("q6", 1, "Từng được thông báo thiếu máu"),
        ("q7", 2, "Từng được thông báo MCV thấp/hồng cầu nhỏ"),
        ("q8", 1, "Từng được thông báo MCH thấp/hồng cầu nhược sắc"),
        ("q10", 2, "Từng được chẩn đoán HbE/hemoglobinopathy khác"),
        ("q11", 1, "Bản thân từng truyền máu nhiều lần/định kỳ"),
        ("q12", 2, "Thiếu máu kéo dài từ nhỏ/tuổi thiếu niên"),
        ("q13", 1, "Mệt mỏi/giảm sức hoạt động"),
        ("q14", 1, "Hoa mắt/chóng mặt không rõ nguyên nhân"),
        ("q15", 1, "Da/niêm nhợt"),
        ("q16", 1, "Vàng da/vàng mắt không rõ nguyên nhân"),
        ("q17", 2, "Từng được ghi nhận lách to/gan lách to"),
        ("q18", 1, "Có tiền sử/biến chứng bệnh huyết học mạn"),
    ]

    for key, weight, label in weighted_items:
        if a.get(key) == "Có":
            score += weight
            reasons.append(label)

    if a.get("q9") == "Đã nghi ngờ":
        score += 2
        reasons.append(
            "Từng có kết quả nghi ngờ Thalassemia/hemoglobinopathy"
        )

    elif a.get("q9") == "Đã xác định mang gen":
        score += 4
        reasons.append(
            "Từng được xác định mang gen"
        )

    return min(
        score,
        ROUND1_MAX_SCORE,
    ), reasons


def round1_category(score):
    if score >= ROUND1_HIGH_THRESHOLD:
        return (
            "CAO",
            "Có đủ yếu tố sàng lọc ban đầu để chuyển sang Vòng 2.",
        )

    if score >= 4:
        return (
            "TRUNG BÌNH",
            "Có một số yếu tố đáng lưu ý nhưng chưa đạt ngưỡng "
            "mở Vòng 2 trong prototype.",
        )

    return (
        "THẤP",
        "Chưa ghi nhận nhiều yếu tố nguy cơ qua bộ câu hỏi ban đầu.",
    )


# ------------------------------------------------------------
# ROUND 2 CBC ANALYSIS
# ------------------------------------------------------------

def calculate_round2_score(
    mcv,
    mch,
    rbc,
    rdw,
):
    score = 0
    reasons = []

    if mcv < 70:
        score += 3
        reasons.append("MCV rất thấp (<70 fL)")
    elif mcv < 75:
        score += 2
        reasons.append("MCV giảm rõ (70–74,9 fL)")
    elif mcv < 80:
        score += 1
        reasons.append("MCV giảm (75–79,9 fL)")

    if mch < 24:
        score += 2
        reasons.append("MCH thấp (<24 pg)")
    elif mch < 27:
        score += 1
        reasons.append("MCH giảm (24–26,9 pg)")

    if mcv < 80:
        if rbc >= 5.5:
            score += 2
            reasons.append("RBC tương đối cao khi MCV thấp")
        elif rbc >= 5.0:
            score += 1
            reasons.append("RBC tương đối cao khi MCV thấp")

    mentzer = (
        mcv / rbc
        if rbc > 0
        else None
    )

    if (
        mentzer is not None
        and mcv < 80
    ):
        if mentzer < 13:
            score += 2
            reasons.append("Mentzer Index <13")
        elif mentzer < 14:
            score += 1
            reasons.append("Mentzer Index 13–13,9")

    if rdw > 15:
        reasons.append(
            "RDW tăng — cần lưu ý thiếu sắt hoặc nguyên nhân microcytosis khác"
        )

    return score, mentzer, reasons


def round2_category(score, mcv):
    if (
        mcv >= 80
        and score <= 2
    ):
        return (
            "THẤP",
            "CBC chưa cho thấy mẫu hình hồng cầu nhỏ rõ.",
        )

    if score <= 3:
        return (
            "THẤP",
            "Nguy cơ sàng lọc từ CBC hiện tại thấp; không loại trừ "
            "hoàn toàn Thalassemia.",
        )

    if score <= 6:
        return (
            "TRUNG BÌNH",
            "Có đặc điểm hồng cầu nhỏ/nhược sắc. Nên đánh giá "
            "thiếu sắt và các nguyên nhân khác.",
        )

    if score <= 9:
        return (
            "CAO",
            "Mẫu hình CBC gợi ý cần đánh giá hemoglobinopathy "
            "bằng HPLC/điện di Hb.",
        )

    return (
        "RẤT CAO",
        "Có nhiều dấu hiệu sàng lọc đáng chú ý; cần xét nghiệm xác nhận.",
    )


def narrative_cbc_advice(
    hb_adjusted,
    mcv,
    mch,
    rbc,
    rdw,
    mentzer,
):
    findings = []
    advice = []

    if mcv < 80:
        findings.append(
            "Có microcytosis (MCV giảm)."
        )

    if mch < 27:
        findings.append(
            "Có xu hướng hồng cầu nhược sắc (MCH giảm)."
        )

    if rdw > 15:
        findings.append(
            "RDW tăng; cần lưu ý thiếu sắt hoặc các nguyên nhân "
            "khác của microcytosis."
        )

    if mcv < 80 and rbc >= 5.0:
        findings.append(
            "RBC tương đối cao trong bối cảnh MCV thấp."
        )

    if mcv < 80 and mentzer < 13:
        findings.append(
            "Mentzer Index <13: mẫu hình sàng lọc nghiêng về "
            "Thalassemia hơn thiếu sắt."
        )

        advice.append(
            "Trao đổi với nhân viên y tế về HPLC/điện di hemoglobin; "
            "tùy trường hợp có thể cần xét nghiệm phân tử."
        )

    elif mcv < 80 and mentzer >= 13:
        findings.append(
            "Mentzer Index ≥13: mẫu hình sàng lọc nghiêng về "
            "thiếu sắt hoặc nguyên nhân microcytosis khác."
        )

        advice.append(
            "Cân nhắc đánh giá tình trạng sắt, đặc biệt Ferritin, "
            "theo chỉ định của nhân viên y tế."
        )

    if hb_adjusted < 8:
        advice.append(
            "Hb sau hiệu chỉnh rất thấp: nên được đánh giá y tế sớm."
        )
    elif hb_adjusted < 10:
        advice.append(
            "Hb sau hiệu chỉnh thấp đáng kể: nên khám và đánh giá nguyên nhân thiếu máu."
        )
    elif hb_adjusted < 12:
        advice.append(
            "Hb sau hiệu chỉnh thấp/giáp ranh tùy nhóm đối tượng; "
            "cần đối chiếu tuổi, giới và khoảng tham chiếu của labo."
        )

    if not findings:
        findings.append(
            "Chưa ghi nhận microcytosis/nhược sắc rõ trên CBC."
        )

    if not advice:
        advice.append(
            "Tiếp tục đối chiếu với khoảng tham chiếu trên phiếu xét nghiệm "
            "và hướng dẫn của nhân viên y tế."
        )

    return findings, advice


# ------------------------------------------------------------
# CURATED MEDICAL FACILITIES BY PROVINCE/CITY
# ------------------------------------------------------------
# Mục tiêu của prototype:
#   1) Ưu tiên bệnh viện hạng I trong chính tỉnh/thành người dùng đang sống.
#   2) Ưu tiên bệnh viện tuyến Trung ương / hạng đặc biệt nếu có tại địa bàn.
#   3) Hiển thị 3–5 cơ sở phù hợp để người dùng chủ động lựa chọn.
#
# Lưu ý:
# - Đây là danh mục điều hướng cho prototype, KHÔNG phải danh sách chỉ định điều trị.
# - Phân hạng/cấp quản lý có thể thay đổi; bản triển khai nghiên cứu chính thức nên
#   đồng bộ định kỳ từ nguồn dữ liệu Bộ Y tế/website chính thức của bệnh viện.
# - Có thêm alias tên tỉnh cũ để app vẫn hoạt động nếu dữ liệu địa giới cũ còn trong DB.

MEDICAL_FACILITIES = {
    "Hà Nội": [
        {"name": "Bệnh viện Bạch Mai", "tier": "Hạng đặc biệt – tuyến Trung ương", "note": "Ưu tiên khi cần đánh giá chuyên sâu huyết học", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Bạch+Mai+Hà+Nội"},
        {"name": "Bệnh viện Hữu nghị Việt Đức", "tier": "Bệnh viện tuyến Trung ương", "note": "Cơ sở chuyên sâu; phù hợp khi cần tuyến trên", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Hữu+nghị+Việt+Đức+Hà+Nội"},
        {"name": "Bệnh viện Trung ương Quân đội 108", "tier": "Hạng đặc biệt – tuyến Trung ương", "note": "Cơ sở tuyến trên với nhiều chuyên khoa sâu", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Trung+ương+Quân+đội+108"},
    ],
    "Hồ Chí Minh": [
        {"name": "Bệnh viện Chợ Rẫy", "tier": "Hạng đặc biệt – tuyến Trung ương", "note": "Ưu tiên khi cần đánh giá chuyên sâu huyết học", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Chợ+Rẫy+Hồ+Chí+Minh"},
        {"name": "Bệnh viện Thống Nhất", "tier": "Hạng I – tuyến Trung ương", "note": "Bệnh viện đa khoa tuyến trên", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Thống+Nhất+Hồ+Chí+Minh"},
        {"name": "Bệnh viện Nhân Dân 115", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến chuyên sâu", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Nhân+Dân+115+Hồ+Chí+Minh"},
        {"name": "Bệnh viện Đại học Y Dược TP. Hồ Chí Minh", "tier": "Hạng I", "note": "Bệnh viện trường đại học tuyến chuyên sâu", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đại+học+Y+Dược+TP+Hồ+Chí+Minh"},
    ],
    "Đà Nẵng": [
        {"name": "Bệnh viện C Đà Nẵng", "tier": "Bệnh viện tuyến Trung ương", "note": "Cơ sở tuyến Trung ương tại Đà Nẵng", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+C+Đà+Nẵng"},
        {"name": "Bệnh viện Đà Nẵng", "tier": "Hạng I", "note": "Bệnh viện đa khoa lớn của thành phố", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đà+Nẵng"},
        {"name": "Bệnh viện Đại học Y – Dược, Đại học Đà Nẵng", "tier": "Bệnh viện trường đại học", "note": "Có thể lựa chọn khi cần đánh giá chuyên khoa", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đại+học+Y+Dược+Đại+học+Đà+Nẵng"},
    ],
    "Thừa Thiên Huế": [],
    "Huế": [
        {"name": "Bệnh viện Trung ương Huế", "tier": "Hạng đặc biệt – tuyến Trung ương", "note": "Ưu tiên khi cần đánh giá chuyên sâu", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Trung+ương+Huế"},
        {"name": "Bệnh viện Trường Đại học Y – Dược Huế", "tier": "Hạng I", "note": "Bệnh viện trường đại học tuyến chuyên sâu", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Trường+Đại+học+Y+Dược+Huế"},
    ],
    "Cần Thơ": [
        {"name": "Bệnh viện Đa khoa Trung ương Cần Thơ", "tier": "Bệnh viện tuyến Trung ương", "note": "Ưu tiên khi cần tuyến chuyên sâu tại ĐBSCL", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+Trung+ương+Cần+Thơ"},
        {"name": "Bệnh viện Đại học Y Dược Cần Thơ", "tier": "Hạng I", "note": "Bệnh viện trường đại học tuyến chuyên sâu", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đại+học+Y+Dược+Cần+Thơ"},
    ],
    "Thái Nguyên": [
        {"name": "Bệnh viện Trung ương Thái Nguyên", "tier": "Hạng đặc biệt – tuyến Trung ương", "note": "Ưu tiên khi cần đánh giá chuyên sâu", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Trung+ương+Thái+Nguyên"},
    ],
    "Hải Phòng": [
        {"name": "Bệnh viện Hữu nghị Việt Tiệp", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến chuyên sâu", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Hữu+nghị+Việt+Tiệp+Hải+Phòng"},
    ],
    "Quảng Ninh": [
        {"name": "Bệnh viện Bãi Cháy", "tier": "Hạng I", "note": "Bệnh viện đa khoa lớn của tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Bãi+Cháy+Quảng+Ninh"},
        {"name": "Bệnh viện Đa khoa tỉnh Quảng Ninh", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+tỉnh+Quảng+Ninh"},
    ],
    "Thanh Hóa": [
        {"name": "Bệnh viện Đa khoa tỉnh Thanh Hóa", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+tỉnh+Thanh+Hóa"},
    ],
    "Nghệ An": [
        {"name": "Bệnh viện Hữu nghị Đa khoa Nghệ An", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Hữu+nghị+Đa+khoa+Nghệ+An"},
    ],
    "Hà Tĩnh": [
        {"name": "Bệnh viện Đa khoa tỉnh Hà Tĩnh", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+tỉnh+Hà+Tĩnh"},
    ],
    "Khánh Hòa": [
        {"name": "Bệnh viện Đa khoa tỉnh Khánh Hòa", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+tỉnh+Khánh+Hòa"},
    ],
    "Đắk Lắk": [
        {"name": "Bệnh viện Đa khoa vùng Tây Nguyên", "tier": "Hạng I", "note": "Cơ sở tuyến chuyên sâu khu vực Tây Nguyên", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+vùng+Tây+Nguyên"},
    ],
    "Lâm Đồng": [
        {"name": "Bệnh viện Đa khoa tỉnh Lâm Đồng", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+tỉnh+Lâm+Đồng"},
    ],
    "Quảng Ngãi": [
        {"name": "Bệnh viện Đa khoa tỉnh Quảng Ngãi", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+tỉnh+Quảng+Ngãi"},
    ],
    "Gia Lai": [
        {"name": "Bệnh viện Đa khoa tỉnh Gia Lai", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+tỉnh+Gia+Lai"},
    ],
    "Bắc Ninh": [
        {"name": "Bệnh viện Đa khoa tỉnh Bắc Ninh", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+tỉnh+Bắc+Ninh"},
    ],
    "Phú Thọ": [
        {"name": "Bệnh viện Đa khoa tỉnh Phú Thọ", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+tỉnh+Phú+Thọ"},
    ],
    "Lào Cai": [
        {"name": "Bệnh viện Đa khoa tỉnh Lào Cai", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+tỉnh+Lào+Cai"},
    ],
    "Hưng Yên": [
        {"name": "Bệnh viện Đa khoa tỉnh Hưng Yên", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+tỉnh+Hưng+Yên"},
    ],
    "Ninh Bình": [
        {"name": "Bệnh viện Đa khoa tỉnh Ninh Bình", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+tỉnh+Ninh+Bình"},
    ],
    "Tây Ninh": [
        {"name": "Bệnh viện Đa khoa tỉnh Tây Ninh", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+tỉnh+Tây+Ninh"},
    ],
    "Đồng Nai": [
        {"name": "Bệnh viện Đa khoa Đồng Nai", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+Đồng+Nai"},
    ],
    "Vĩnh Long": [
        {"name": "Bệnh viện Đa khoa tỉnh Vĩnh Long", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+tỉnh+Vĩnh+Long"},
    ],
    "Đồng Tháp": [
        {"name": "Bệnh viện Đa khoa Đồng Tháp", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+Đồng+Tháp"},
    ],
    "An Giang": [
        {"name": "Bệnh viện Đa khoa khu vực An Giang", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh/khu vực", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+khu+vực+An+Giang"},
    ],
    "Cà Mau": [
        {"name": "Bệnh viện Đa khoa tỉnh Cà Mau", "tier": "Hạng I", "note": "Bệnh viện đa khoa tuyến tỉnh", "maps": "https://www.google.com/maps/search/?api=1&query=Bệnh+viện+Đa+khoa+tỉnh+Cà+Mau"},
    ],
}

# Tên tỉnh cũ → tên tỉnh/thành dùng để tra danh mục cơ sở trong prototype.
PROVINCE_FACILITY_ALIASES = {
    "Thừa Thiên Huế": "Huế",
    "Quảng Nam": "Đà Nẵng",
    "Bình Định": "Gia Lai",
    "Ninh Thuận": "Khánh Hòa",
    "Phú Yên": "Đắk Lắk",
    "Đắk Nông": "Lâm Đồng",
    "Bình Thuận": "Lâm Đồng",
    "Kon Tum": "Quảng Ngãi",
    "Yên Bái": "Lào Cai",
    "Bắc Kạn": "Thái Nguyên",
    "Hòa Bình": "Phú Thọ",
    "Vĩnh Phúc": "Phú Thọ",
    "Hà Nam": "Ninh Bình",
    "Nam Định": "Ninh Bình",
    "Quảng Bình": "Quảng Trị",
    "Bà Rịa - Vũng Tàu": "Hồ Chí Minh",
    "Bình Dương": "Hồ Chí Minh",
    "Long An": "Tây Ninh",
    "Tiền Giang": "Đồng Tháp",
    "Bến Tre": "Vĩnh Long",
    "Trà Vinh": "Vĩnh Long",
    "Sóc Trăng": "Cần Thơ",
    "Hậu Giang": "Cần Thơ",
    "Kiên Giang": "An Giang",
}


def recommended_facilities(province):
    """Trả về 3–5 cơ sở ưu tiên theo tỉnh/thành hiện tại."""
    province_key = PROVINCE_FACILITY_ALIASES.get(province, province)
    facilities = MEDICAL_FACILITIES.get(province_key, [])
    return facilities[:5]


# ------------------------------------------------------------
# GOOGLE PLACES
# ------------------------------------------------------------

@st.cache_data(ttl=3600)
def nearby_medical(
    lat,
    lng,
    api_key,
):
    if not api_key:
        return []

    url = (
        "https://places.googleapis.com/v1/places:searchNearby"
    )

    headers = {
        "Content-Type": "application/json",
        "X-Goog-Api-Key": api_key,
        "X-Goog-FieldMask": (
            "places.displayName,"
            "places.formattedAddress,"
            "places.location,"
            "places.googleMapsUri,"
            "places.rating,"
            "places.userRatingCount"
        ),
    }

    payload = {
        "includedTypes": [
            "hospital",
            "medical_center",
            "medical_clinic",
            "medical_lab",
        ],
        "maxResultCount": 20,
        "rankPreference": "DISTANCE",
        "locationRestriction": {
            "circle": {
                "center": {
                    "latitude": lat,
                    "longitude": lng,
                },
                "radius": 50000.0,
            }
        },
    }

    try:
        response = requests.post(
            url,
            headers=headers,
            json=payload,
            timeout=15,
        )
        response.raise_for_status()
        return response.json().get(
            "places",
            [],
        )
    except Exception:
        return []


def haversine_km(
    lat1,
    lon1,
    lat2,
    lon2,
):
    radius = 6371.0

    p1 = math.radians(lat1)
    p2 = math.radians(lat2)

    dp = math.radians(
        lat2 - lat1
    )

    dl = math.radians(
        lon2 - lon1
    )

    a = (
        math.sin(dp / 2) ** 2
        + math.cos(p1)
        * math.cos(p2)
        * math.sin(dl / 2) ** 2
    )

    return (
        2
        * radius
        * math.asin(
            math.sqrt(a)
        )
    )


def rank_facilities(
    places,
    lat,
    lng,
):
    rows = []

    for place in places:
        loc = place.get(
            "location",
            {},
        )

        lat2 = loc.get(
            "latitude"
        )
        lng2 = loc.get(
            "longitude"
        )

        if (
            lat2 is None
            or lng2 is None
        ):
            continue

        rows.append(
            {
                "name": place.get(
                    "displayName",
                    {},
                ).get(
                    "text",
                    "Cơ sở y tế",
                ),
                "address": place.get(
                    "formattedAddress",
                    "Chưa có địa chỉ",
                ),
                "distance": haversine_km(
                    lat,
                    lng,
                    lat2,
                    lng2,
                ),
                "rating": place.get(
                    "rating"
                ),
                "rating_count": place.get(
                    "userRatingCount"
                ),
                "maps_uri": place.get(
                    "googleMapsUri"
                ),
            }
        )

    rows.sort(
        key=lambda x: x["distance"]
    )

    return rows[:5]


# ------------------------------------------------------------
# LOW RISK PANEL
# ------------------------------------------------------------

def low_risk_panel():

    st.success(
        "🟢 **NGUY CƠ SÀNG LỌC BAN ĐẦU: THẤP**"
    )

    st.markdown(
        """
### 📅 Theo dõi sức khỏe

Kết quả Vòng 1 hiện chưa cho thấy nhiều yếu tố nguy cơ rõ ràng.

Bạn nên tiếp tục **theo dõi tình trạng sức khỏe và khám định kỳ
theo hướng dẫn của cơ sở y tế**.

Trong prototype, hệ thống đặt mốc xem xét lại khoảng **30 ngày**.
Đây là mốc theo dõi của ứng dụng, không phải chỉ định bắt buộc
mọi người nguy cơ thấp phải khám hàng tháng.

### 🩸 Khi nào nên kiểm tra CBC lại?

Nếu xuất hiện mệt mỏi kéo dài, da/niêm nhợt, chóng mặt, vàng da/vàng mắt
hoặc phiếu công thức máu có bất thường, hãy đưa kết quả cho nhân viên y tế.
Bạn có thể nhập CBC bên dưới để hệ thống **sàng lọc lại**.
"""
    )

    reminder_date = (
        date.today()
        + timedelta(
            days=FOLLOWUP_DAYS
        )
    )

    st.info(
        f"🗓️ Mốc nhắc prototype: "
        f"**{reminder_date.strftime('%d/%m/%Y')}**"
    )

    with st.expander(
        "🩸 Nhập CBC nếu lần xét nghiệm sau có bất thường",
        expanded=False,
    ):

        a, b, c, d = st.columns(4)

        with a:
            hb_unit = st.selectbox(
                "Đơn vị Hb",
                ["g/dL", "g/L"],
                key="low_cbc_hb_unit",
            )
            hb_raw = st.number_input(
                "Hb",
                min_value=3.0 if hb_unit == "g/dL" else 30.0,
                max_value=25.0 if hb_unit == "g/dL" else 250.0,
                value=13.0 if hb_unit == "g/dL" else 130.0,
                step=0.1,
                key="low_cbc_hb",
            )

        with b:
            mcv = st.number_input(
                "MCV (fL)",
                30.0,
                150.0,
                85.0,
                0.1,
                key="low_cbc_mcv",
            )

        with c:
            mch = st.number_input(
                "MCH (pg)",
                10.0,
                50.0,
                29.0,
                0.1,
                key="low_cbc_mch",
            )

        with d:
            rbc_unit = st.selectbox(
                "Đơn vị RBC",
                ["T/L", "10^12/L", "10^6/µL"],
                key="low_cbc_rbc_unit",
            )
            rbc_raw = st.number_input(
                "RBC",
                1.0,
                10.0,
                4.8,
                0.1,
                key="low_cbc_rbc",
            )

        rdw = st.number_input(
            "RDW-CV (%)",
            5.0,
            40.0,
            13.0,
            0.1,
            key="low_cbc_rdw",
        )

        if st.button(
            "🔎 PHÂN TÍCH CBC BẤT THƯỜNG",
            key="low_cbc_analyze",
        ):

            hb = hb_to_g_dl(
                hb_raw,
                hb_unit,
            )

            rbc = rbc_to_t_l(
                rbc_raw,
                rbc_unit,
            )

            score, mentzer, reasons = (
                calculate_round2_score(
                    mcv,
                    mch,
                    rbc,
                    rdw,
                )
            )

            category, conclusion = (
                round2_category(
                    score,
                    mcv,
                )
            )

            if category == "THẤP":
                st.success(
                    f"🟢 {category} — {conclusion}"
                )
            elif category == "TRUNG BÌNH":
                st.warning(
                    f"🟡 {category} — {conclusion}"
                )
            else:
                st.error(
                    f"🟠 {category} — {conclusion}"
                )

            st.metric(
                "Mentzer Index",
                f"{mentzer:.2f}",
            )

            findings, advice = narrative_cbc_advice(
                hb,
                mcv,
                mch,
                rbc,
                rdw,
                mentzer,
            )

            st.markdown(
                "### 🧠 Phân tích sơ bộ"
            )

            for finding in findings:
                st.write(
                    f"• {finding}"
                )

            st.markdown(
                "### 💡 Khuyến nghị"
            )

            for item in advice:
                st.write(
                    f"→ {item}"
                )


# ------------------------------------------------------------
# WORD REPORT
# ------------------------------------------------------------

def make_word(
    patient,
    r1_score,
    r1_category,
    r1_reasons,
    r2=None,
):
    doc = Document()

    doc.add_heading(
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        level=3,
    )

    doc.add_heading(
        "PHIẾU SÀNG LỌC VÀ PHÂN TẦNG NGUY CƠ THALASSEMIA",
        level=1,
    )

    doc.add_paragraph(
        "Công cụ hỗ trợ sàng lọc; không thay thế chẩn đoán "
        "hoặc chỉ định của nhân viên y tế."
    )

    doc.add_heading(
        "I. THÔNG TIN BỆNH NHÂN",
        level=2,
    )

    for label, value in [
        ("Họ và tên", patient["full_name"]),
        ("Ngày sinh", patient["birth_date"]),
        ("Tuổi", patient["age"]),
        ("Giới tính", patient["gender"]),
        ("Số điện thoại", patient["phone"]),
        ("Địa chỉ hiện tại", patient["current_address"]),
        ("Tỉnh/thành", patient["province"]),
        ("Phường/xã/đặc khu", patient["commune"]),
    ]:
        doc.add_paragraph(
            f"{label}: {value}"
        )

    doc.add_heading(
        "II. VÒNG 1",
        level=2,
    )

    doc.add_paragraph(
        f"Điểm: {r1_score}/{ROUND1_MAX_SCORE}"
    )

    doc.add_paragraph(
        f"Mức nguy cơ: {r1_category}"
    )

    for item in r1_reasons:
        doc.add_paragraph(
            f"- {item}"
        )

    if r2:

        doc.add_heading(
            "III. VÒNG 2",
            level=2,
        )

        doc.add_paragraph(
            f"Khoảng độ cao: {r2['altitude_choice']}"
        )

        doc.add_paragraph(
            f"Hiệu chỉnh Hb: -{r2['adjustment']:.1f} g/dL"
        )

        doc.add_paragraph(
            f"Hb thực đo sau quy đổi: {r2['hb']:.1f} g/dL"
        )

        doc.add_paragraph(
            f"Hb sau hiệu chỉnh: {r2['hb_adjusted']:.1f} g/dL"
        )

        doc.add_paragraph(
            f"MCV: {r2['mcv']:.1f} fL"
        )

        doc.add_paragraph(
            f"MCH: {r2['mch']:.1f} pg"
        )

        doc.add_paragraph(
            f"RBC: {r2['rbc']:.2f} T/L"
        )

        doc.add_paragraph(
            f"RDW-CV: {r2['rdw']:.1f}%"
        )

        doc.add_paragraph(
            f"Mentzer Index: {r2['mentzer']:.2f}"
        )

        doc.add_paragraph(
            f"Điểm CBC prototype: {r2['score']}"
        )

        doc.add_paragraph(
            f"Mức nguy cơ Vòng 2: {r2['category']}"
        )

        doc.add_paragraph(
            f"Nhận định: {r2['conclusion']}"
        )

        doc.add_heading(
            "IV. PHÂN TÍCH SƠ BỘ",
            level=2,
        )

        for finding in r2["findings"]:
            doc.add_paragraph(
                f"- {finding}"
            )

        doc.add_heading(
            "V. KHUYẾN NGHỊ",
            level=2,
        )

        for advice in r2["advice"]:
            doc.add_paragraph(
                f"- {advice}"
            )

    else:

        doc.add_heading(
            "III. THEO DÕI",
            level=2,
        )

        doc.add_paragraph(
            "Vòng 2 chưa được mở. Tiếp tục theo dõi sức khỏe "
            "và đánh giá lại nếu xuất hiện bất thường."
        )

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)

    return out


# ============================================================
# HEADER
# ============================================================

st.title(
    "🩸 HỆ THỐNG SÀNG LỌC VÀ PHÂN TẦNG NGUY CƠ THALASSEMIA"
)

st.write(
    "Hồ sơ bệnh nhân → Vòng 1 (20 câu) → nguy cơ cao → "
    "Vòng 2 (độ cao + CBC) → phân tích sơ bộ → khuyến nghị."
)

st.info(
    "🎯 Mục tiêu: hỗ trợ sàng lọc ban đầu tại tuyến cơ sở "
    "và điều hướng người có nguy cơ tới dịch vụ phù hợp."
)


# ============================================================
# ACCESS / SIDEBAR
# ============================================================

with st.sidebar:
    st.header("⚙️ Trạng thái kỹ thuật")

    if GOOGLE_API_KEY:
        st.success("Google Places: đã cấu hình")
    else:
        st.caption("Google Places: chưa cấu hình API key — hệ thống vẫn có danh mục cơ sở ưu tiên theo tỉnh/thành.")

    st.divider()
    st.caption(
        "👤 Hồ sơ → 🟦 Vòng 1 → 🔴 Nguy cơ cao → "
        "🟧 Vòng 2 → 🧠 Phân tích → 🏥 Gợi ý cơ sở"
    )

render_auth_sidebar()

# ============================================================
# ADMIN CONSOLE
# ============================================================

auth_user = current_auth_user()

if auth_user and st.session_state.get("auth_page") == "🛡️ Quản trị hệ thống":
    render_admin_console(auth_user)
    st.stop()

operator_mode = auth_user is not None
operator_username = auth_user["username"] if auth_user else None

if operator_mode:
    st.success(
        f"👩‍🔬 **Chế độ nhập giúp người tham gia** — thao tác đang được ghi nhận bởi **{auth_user['full_name']}** (#{auth_user['username']})."
    )

# ============================================================
# PATIENT PROFILE
# ============================================================

st.header(
    "👤 THÔNG TIN BỆNH NHÂN"
)
st.caption(
    "Khu vực này dùng để nhập hồ sơ người tham gia. "
    + (
        "Bạn đang ở chế độ **nhập giúp người tham gia**; dữ liệu sẽ ghi nhận tài khoản nhân sự đang đăng nhập."
        if operator_mode
        else "Người tham gia có thể tự nhập thông tin của chính mình."
    )
    + " Danh sách hồ sơ của nhiều người chỉ được hiển thị trong khu vực quản trị có kiểm soát truy cập."
)

with st.container(border=True):

    p1, p2, p3 = st.columns(3)

    with p1:

        full_name = st.text_input(
            "Họ và tên *",
            placeholder="Nguyễn Văn A",
        )

        birth_date = st.date_input(
            "Ngày sinh *",
            value=date(
                2000,
                1,
                1,
            ),
            min_value=date(
                1900,
                1,
                1,
            ),
            max_value=date.today(),
            format="DD/MM/YYYY",
        )

    with p2:

        phone_raw = st.text_input(
            "Số điện thoại *",
            placeholder="09xxxxxxxx",
        )

        gender = st.selectbox(
            "Giới tính *",
            [
                "Nam",
                "Nữ",
                "Khác",
            ],
        )

    with p3:

        current_address = st.text_input(
            "Địa chỉ hiện tại *",
            placeholder="Số nhà/thôn/tổ/đường",
        )

    st.subheader(
        "📍 Địa giới hành chính hiện tại"
    )

    try:
        admin_hierarchy = load_admin_hierarchy()
        admin_data_ok = True
    except Exception as exc:
        admin_hierarchy = {}
        admin_data_ok = False
        st.error(
            "Không tải được danh mục tỉnh/thành và phường/xã hiện hành. "
            "Vui lòng tải lại trang hoặc thử lại sau."
        )
        st.caption(
            f"Nguồn dữ liệu: {ADMIN_DATA_SOURCE_URL}"
        )
        st.code(str(exc))

    if admin_data_ok:
        provinces_list = sorted(
            admin_hierarchy.keys(),
            key=str.casefold,
        )

        selected_province = st.selectbox(
            "Tỉnh / thành phố *",
            ["— Chọn tỉnh/thành —"] + provinces_list,
            key="profile_province",
        )

        if selected_province == "— Chọn tỉnh/thành —":
            commune_value = ""
            st.selectbox(
                "Phường / xã / đặc khu *",
                ["— Chọn tỉnh/thành trước —"],
                disabled=True,
                key="profile_commune_disabled",
            )
        else:
            available_communes = admin_hierarchy[selected_province]
            commune_choice = st.selectbox(
                "Phường / xã / đặc khu *",
                ["— Chọn phường/xã/đặc khu —"] + available_communes,
                key="profile_commune",
            )
            commune_value = (
                ""
                if commune_choice == "— Chọn phường/xã/đặc khu —"
                else commune_choice
            )

        st.caption(
            "Danh mục địa giới được tải theo cấu trúc 2 cấp hiện hành; "
            "không nhập tay tên phường/xã để tránh sai địa danh."
        )
    else:
        selected_province = ""
        commune_value = ""

    if selected_province and commune_value:
        st.success(
            f"📍 Đã chọn: **{commune_value}, {selected_province}**"
        )

    st.markdown("### 🔐 Đồng ý tham gia sàng lọc và nghiên cứu")
    if operator_mode:
        st.info(
            "Bạn đang nhập giúp người tham gia. Chỉ tiếp tục khi người tham gia đã được giải thích "
            "nội dung, đồng ý cho nghiên cứu sinh sử dụng dữ liệu theo mục đích nghiên cứu/sàng lọc, "
            "và bạn có cơ sở hợp lý để ghi nhận sự đồng ý đó."
        )
    else:
        st.info(
            "Để tiếp tục, người tham gia cần đọc và đồng ý với nội dung dưới đây. "
            "Nếu không đồng ý, hệ thống sẽ **không thực hiện sàng lọc và không lưu hồ sơ/thông tin sức khỏe**."
        )
    with st.container(border=True):
        st.markdown(
            "**Tôi đồng ý cho nghiên cứu sinh sử dụng thông tin cá nhân, "
            "thông tin khảo sát và thông tin sức khỏe/xét nghiệm do tôi cung cấp "
            "cho mục đích sàng lọc cộng đồng Thalassemia và nghiên cứu khoa học.**"
        )
        st.markdown(
            "Tôi hiểu rằng việc tham gia là tự nguyện; kết quả của hệ thống chỉ có "
            "tính chất sàng lọc, không thay thế chẩn đoán của cơ sở y tế; dữ liệu "
            "được lưu phục vụ mục đích nêu trên theo phiên bản chấp thuận của nghiên cứu. "
            "Tôi có thể dừng tham gia bằng cách không tiếp tục sử dụng hệ thống."
        )
        research_consent = st.checkbox(
            "✅ Tôi đã đọc, hiểu và đồng ý tham gia." if not operator_mode
            else "✅ Tôi xác nhận người tham gia đã đọc/được giải thích và đã đồng ý.",
            key="research_consent",
        )
        if operator_mode:
            st.caption(
                f"Người nhập: {auth_user['full_name']} ({auth_user['username']})"
            )
        st.caption(
            f"Phiên bản nội dung chấp thuận: {CONSENT_VERSION}"
        )


phone = normalize_phone(
    phone_raw
)

if (
    research_consent
    and phone
    and valid_vietnam_phone(phone)
    and phone_exists(phone)
):

    st.warning(
        "📌 Số điện thoại này đã tồn tại. "
        "Lưu lại sẽ **ghi đè bằng lần nhập sau cùng**."
    )


if st.button(
    "💾 LƯU / CẬP NHẬT HỒ SƠ",
    type="secondary",
    disabled=not (research_consent and admin_data_ok),
):

    if not research_consent:

        st.error(
            "Bạn cần đồng ý tham gia sàng lọc và nghiên cứu trước khi tiếp tục."
        )

    elif not admin_data_ok:

        st.error(
            "Chưa tải được danh mục địa giới hiện hành nên chưa thể lưu hồ sơ."
        )

    elif not full_name.strip():

        st.error(
            "Vui lòng nhập họ và tên."
        )

    elif not valid_vietnam_phone(phone):

        st.error(
            "Số điện thoại phải là số Việt Nam 10 chữ số."
        )

    elif not current_address.strip():

        st.error(
            "Vui lòng nhập địa chỉ hiện tại."
        )

    elif not selected_province:

        st.error(
            "Vui lòng chọn tỉnh/thành phố."
        )

    elif not commune_value:

        st.error(
            "Vui lòng chọn phường/xã/đặc khu."
        )

    else:

        profile = {
            "phone": phone,
            "full_name": full_name.strip(),
            "birth_date": birth_date.isoformat(),
            "gender": gender,
            "current_address": current_address.strip(),
            "province": selected_province,
            "commune": commune_value,
            "age": calculate_age(
                birth_date
            ),
            "consent_at": datetime.now().isoformat(timespec="seconds"),
            "entry_mode": "assisted" if operator_mode else "self",
            "entered_by_username": operator_username,
        }

        action = upsert_patient(
            profile
        )

        st.session_state[
            "patient_profile"
        ] = profile

        reset_results()

        if action == "updated":
            st.success(
                "✅ Đã cập nhật hồ sơ bằng lần nhập sau cùng."
            )
        else:
            st.success(
                "✅ Đã tạo hồ sơ bệnh nhân."
            )


patient = st.session_state.get(
    "patient_profile"
)

if not st.session_state.get("research_consent", False):
    st.warning(
        "🔒 Bạn chưa đồng ý tham gia. Hệ thống không mở Vòng 1 và không xử lý/lưu dữ liệu sàng lọc."
    )
    st.stop()

if not patient:

    st.info(
        "👆 Hãy lưu hồ sơ bệnh nhân trước khi bắt đầu Vòng 1."
    )

    st.stop()


# ============================================================
# PROFILE SUMMARY
# ============================================================

st.success(
    f"✅ **{patient['full_name']}** · "
    f"{patient['age']} tuổi · "
    f"{patient['phone']} · "
    f"{patient['commune']}, {patient['province']}"
)


# ============================================================
# ROUND 1 QUESTIONS
# ============================================================

st.divider()

st.header(
    "🟦 VÒNG 1 — 20 CÂU HỎI SÀNG LỌC"
)

st.info(
    "Q1–Q18 phục vụ sàng lọc nguy cơ. "
    "**Q19–Q20 không cộng điểm** vì chỉ phản ánh khả năng tiếp cận xét nghiệm."
)


with st.container(border=True):

    st.subheader(
        "A. Tiền sử gia đình"
    )

    q1 = st.radio(
        "1. Trong gia đình/dòng họ có người từng được chẩn đoán Thalassemia không?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )

    q2 = st.radio(
        "2. Trong gia đình/dòng họ có người từng được thông báo mang gen Thalassemia/hemoglobinopathy không?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )

    q3 = st.radio(
        "3. Cha hoặc mẹ bạn có từng được xét nghiệm Thalassemia/hemoglobinopathy không?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )

    q4 = st.radio(
        "4. Anh/chị/em ruột có từng được chẩn đoán thiếu máu hoặc hồng cầu nhỏ không?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )

    q5 = st.radio(
        "5. Gia đình có trẻ từng phải truyền máu nhiều lần hoặc định kỳ không?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )


with st.container(border=True):

    st.subheader(
        "B. Tiền sử bản thân"
    )

    q6 = st.radio(
        "6. Bạn từng được nhân viên y tế thông báo bị thiếu máu chưa?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )

    q7 = st.radio(
        "7. Bạn từng được thông báo MCV thấp/hồng cầu nhỏ chưa?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )

    q8 = st.radio(
        "8. Bạn từng được thông báo MCH thấp/hồng cầu nhược sắc chưa?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )

    q9 = st.selectbox(
        "9. Bạn từng xét nghiệm Thalassemia/hemoglobinopathy chưa?",
        [
            "Chưa xét nghiệm",
            "Đã xét nghiệm, bình thường",
            "Đã nghi ngờ",
            "Đã xác định mang gen",
            "Không nhớ",
        ],
    )

    q10 = st.radio(
        "10. Bạn từng được chẩn đoán HbE hoặc hemoglobinopathy khác chưa?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )

    q11 = st.radio(
        "11. Bản thân từng truyền máu nhiều lần hoặc định kỳ chưa?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )

    q12 = st.radio(
        "12. Bạn có tiền sử thiếu máu kéo dài từ nhỏ hoặc từ tuổi thiếu niên không?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )


with st.container(border=True):

    st.subheader(
        "C. Dấu hiệu hỗ trợ"
    )

    q13 = st.radio(
        "13. Bạn có thường xuyên mệt mỏi hoặc giảm khả năng hoạt động không?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )

    q14 = st.radio(
        "14. Bạn có thường xuyên hoa mắt/chóng mặt không rõ nguyên nhân không?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )

    q15 = st.radio(
        "15. Bạn từng được nhận xét da hoặc niêm mạc nhợt hơn bình thường chưa?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )

    q16 = st.radio(
        "16. Bạn từng có vàng da/vàng mắt không rõ nguyên nhân chưa?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )

    q17 = st.radio(
        "17. Bạn từng được bác sĩ ghi nhận lách to hoặc gan lách to chưa?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )

    q18 = st.radio(
        "18. Bạn từng được bác sĩ lưu ý có biến chứng liên quan bệnh huyết học mạn chưa?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )


with st.container(border=True):

    st.subheader(
        "D. Khả năng tiếp cận xét nghiệm — KHÔNG TÍNH ĐIỂM"
    )

    q19 = st.radio(
        "19. Bạn hiện có CBC trong vòng 6–12 tháng gần đây không?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )

    q20 = st.radio(
        "20. Bạn có gặp khó khăn khi đến cơ sở có xét nghiệm chuyên sâu "
        "do khoảng cách, chi phí hoặc thời gian di chuyển không?",
        ["Không", "Có", "Không biết"],
        horizontal=True,
    )

    st.info(
        "Q19–Q20 được lưu để hỗ trợ điều hướng y tế, **không ảnh hưởng "
        "đến điểm nguy cơ Thalassemia**."
    )


if st.button(
    "🔎 ĐÁNH GIÁ VÒNG 1",
    type="primary",
    use_container_width=True,
):

    answers = {
        "q1": q1,
        "q2": q2,
        "q3": q3,
        "q4": q4,
        "q5": q5,
        "q6": q6,
        "q7": q7,
        "q8": q8,
        "q9": q9,
        "q10": q10,
        "q11": q11,
        "q12": q12,
        "q13": q13,
        "q14": q14,
        "q15": q15,
        "q16": q16,
        "q17": q17,
        "q18": q18,
        "q19": q19,
        "q20": q20,
    }

    score1, reasons1 = (
        calculate_round1_score(
            answers
        )
    )

    category1, conclusion1 = (
        round1_category(
            score1
        )
    )

    st.session_state[
        "round1_score"
    ] = score1

    st.session_state[
        "round1_reasons"
    ] = reasons1

    # Lưu một lượt sàng lọc hoàn chỉnh Vòng 1. Nếu có tài khoản đăng nhập,
    # ghi nhận người nhập là nhân sự/ quản trị viên đang thực hiện thao tác.
    st.session_state["screening_id"] = create_screening_record(
        patient=patient,
        answers=answers,
        score1=score1,
        category1=category1,
        conclusion1=conclusion1,
        reasons1=reasons1,
        entry_mode="assisted" if operator_mode else "self",
        entered_by_username=operator_username,
    )

    st.session_state[
        "round1_category"
    ] = category1

    st.session_state[
        "round1_conclusion"
    ] = conclusion1

    st.session_state[
        "round1_completed"
    ] = True

    # Vòng 1 mới -> xóa Vòng 2 cũ.
    for key in list(
        st.session_state.keys()
    ):
        if (
            key.startswith("round2_")
            or key.startswith("google_")
        ):
            del st.session_state[key]


# ============================================================
# ROUND 1 RESULT
# ============================================================

if st.session_state.get(
    "round1_completed",
    False,
):

    st.subheader(
        "📋 KẾT QUẢ VÒNG 1"
    )

    score1 = st.session_state[
        "round1_score"
    ]

    category1 = st.session_state[
        "round1_category"
    ]

    conclusion1 = st.session_state[
        "round1_conclusion"
    ]

    reasons1 = st.session_state[
        "round1_reasons"
    ]

    a1, a2 = st.columns(2)

    with a1:
        st.metric(
            "Điểm Vòng 1",
            f"{score1}/{ROUND1_MAX_SCORE}",
        )

    with a2:
        st.metric(
            "Ngưỡng mở Vòng 2",
            f"≥{ROUND1_HIGH_THRESHOLD}",
        )

    if category1 == "CAO":

        st.error(
            f"🔴 **NGUY CƠ: CAO**\n\n"
            f"{conclusion1}"
        )

        if reasons1:
            with st.expander(
                "🔎 Các yếu tố đáng chú ý",
                expanded=True,
            ):
                for item in reasons1:
                    st.write(
                        f"• {item}"
                    )

        st.session_state[
            "round2_unlocked"
        ] = True

        st.success(
            "✅ Vòng 2 đã được mở."
        )

    elif category1 == "TRUNG BÌNH":

        st.warning(
            f"🟡 **NGUY CƠ: TRUNG BÌNH**\n\n"
            f"{conclusion1}"
        )

        if reasons1:
            with st.expander(
                "🔎 Các yếu tố đáng chú ý",
                expanded=False,
            ):
                for item in reasons1:
                    st.write(
                        f"• {item}"
                    )

        st.info(
            "Prototype chưa tự động mở Vòng 2 ở mức trung bình."
        )

    else:

        low_risk_panel()


# ============================================================
# ROUND 2
# ============================================================

if st.session_state.get(
    "round2_unlocked",
    False,
):

    st.divider()

    st.header(
        "🟧 VÒNG 2 — ĐỘ CAO + CBC"
    )

    # --------------------------------------------------------
    # LOCATION + ALTITUDE
    # --------------------------------------------------------

    with st.container(border=True):

        st.subheader(
            "1. Nơi cư trú và độ cao"
        )

        st.write(
            f"**Tỉnh/thành:** {patient['province']}"
        )

        st.write(
            f"**Phường/xã/đặc khu:** {patient['commune']}"
        )

        st.write(
            f"**Địa chỉ hiện tại:** {patient['current_address']}"
        )

        altitude_choice, altitude_adjustment = (
            altitude_selector(
                "round2"
            )
        )

        st.session_state[
            "round2_altitude_choice"
        ] = altitude_choice

        st.session_state[
            "round2_adjustment"
        ] = altitude_adjustment

    # --------------------------------------------------------
    # CBC
    # --------------------------------------------------------

    with st.container(border=True):

        st.subheader(
            "2. Nhập Công thức máu"
        )

        c1, c2, c3, c4, c5 = st.columns(5)

        with c1:

            hb_unit = st.selectbox(
                "Đơn vị Hb",
                ["g/dL", "g/L"],
                key="round2_hb_unit",
            )

            hb_raw = st.number_input(
                "Hb",
                min_value=3.0
                if hb_unit == "g/dL"
                else 30.0,
                max_value=25.0
                if hb_unit == "g/dL"
                else 250.0,
                value=13.0
                if hb_unit == "g/dL"
                else 130.0,
                step=0.1,
                key="round2_hb_raw",
            )

        with c2:

            mcv = st.number_input(
                "MCV (fL)",
                30.0,
                150.0,
                85.0,
                0.1,
                key="round2_mcv",
            )

        with c3:

            mch = st.number_input(
                "MCH (pg)",
                10.0,
                50.0,
                29.0,
                0.1,
                key="round2_mch",
            )

        with c4:

            rbc_unit = st.selectbox(
                "Đơn vị RBC",
                [
                    "T/L",
                    "10^12/L",
                    "10^6/µL",
                ],
                key="round2_rbc_unit",
            )

            rbc_raw = st.number_input(
                "RBC",
                1.0,
                10.0,
                4.8,
                0.1,
                key="round2_rbc_raw",
            )

        with c5:

            rdw = st.number_input(
                "RDW-CV (%)",
                5.0,
                40.0,
                13.0,
                0.1,
                key="round2_rdw",
            )

        st.caption(
            "Hệ thống tự chuyển Hb → g/dL và RBC → T/L trước khi tính."
        )


    # --------------------------------------------------------
    # ANALYZE
    # --------------------------------------------------------

    if st.button(
        "🩸 PHÂN TÍCH VÒNG 2",
        type="primary",
        use_container_width=True,
    ):

        if rbc_raw <= 0:

            st.error(
                "RBC phải lớn hơn 0."
            )

        else:

            hb = hb_to_g_dl(
                hb_raw,
                hb_unit,
            )

            rbc = rbc_to_t_l(
                rbc_raw,
                rbc_unit,
            )

            adjustment = st.session_state[
                "round2_adjustment"
            ]

            hb_adjusted = (
                hb - adjustment
            )

            score2, mentzer, reasons2 = (
                calculate_round2_score(
                    mcv,
                    mch,
                    rbc,
                    rdw,
                )
            )

            category2, conclusion2 = (
                round2_category(
                    score2,
                    mcv,
                )
            )

            findings, advice = (
                narrative_cbc_advice(
                    hb_adjusted,
                    mcv,
                    mch,
                    rbc,
                    rdw,
                    mentzer,
                )
            )

            st.session_state[
                "round2_score"
            ] = score2

            st.session_state[
                "round2_mentzer"
            ] = mentzer

            st.session_state[
                "round2_category"
            ] = category2

            st.session_state[
                "round2_conclusion"
            ] = conclusion2

            st.session_state[
                "round2_reasons"
            ] = reasons2

            st.session_state[
                "round2_hb"
            ] = hb

            st.session_state[
                "round2_hb_adjusted"
            ] = hb_adjusted

            st.session_state[
                "round2_mcv_result"
            ] = mcv

            st.session_state[
                "round2_mch_result"
            ] = mch

            st.session_state[
                "round2_rbc_result"
            ] = rbc

            st.session_state[
                "round2_rdw_result"
            ] = rdw

            st.session_state[
                "round2_findings"
            ] = findings

            st.session_state[
                "round2_advice"
            ] = advice

            update_screening_round2(
                st.session_state.get("screening_id"),
                {
                    "altitude_choice": st.session_state["round2_altitude_choice"],
                    "adjustment": adjustment,
                    "hb": hb,
                    "hb_adjusted": hb_adjusted,
                    "mcv": mcv,
                    "mch": mch,
                    "rbc": rbc,
                    "rdw": rdw,
                    "mentzer": mentzer,
                    "score": score2,
                    "category": category2,
                    "conclusion": conclusion2,
                    "reasons": reasons2,
                    "findings": findings,
                    "advice": advice,
                },
            )

            st.session_state[
                "round2_completed"
            ] = True


    # --------------------------------------------------------
    # RESULTS
    # --------------------------------------------------------

    if st.session_state.get(
        "round2_completed",
        False,
    ):

        st.subheader(
            "📊 KẾT QUẢ VÒNG 2"
        )

        a, b, c, d = st.columns(4)

        with a:
            st.metric(
                "Mentzer Index",
                f"{st.session_state['round2_mentzer']:.2f}",
            )

        with b:
            st.metric(
                "Điểm CBC",
                str(
                    st.session_state[
                        "round2_score"
                    ]
                ),
            )

        with c:
            st.metric(
                "Hb thực đo",
                f"{st.session_state['round2_hb']:.1f} g/dL",
            )

        with d:
            st.metric(
                "Hb sau hiệu chỉnh",
                f"{st.session_state['round2_hb_adjusted']:.1f} g/dL",
            )

        category2 = st.session_state[
            "round2_category"
        ]

        conclusion2 = st.session_state[
            "round2_conclusion"
        ]

        if category2 == "THẤP":

            st.success(
                f"🟢 **NGUY CƠ VÒNG 2: THẤP**\n\n"
                f"{conclusion2}"
            )

        elif category2 == "TRUNG BÌNH":

            st.warning(
                f"🟡 **NGUY CƠ VÒNG 2: TRUNG BÌNH**\n\n"
                f"{conclusion2}"
            )

        else:

            st.error(
                f"🔴 **NGUY CƠ VÒNG 2: {category2}**\n\n"
                f"{conclusion2}"
            )

        with st.expander(
            "🔎 Các yếu tố từ CBC",
            expanded=True,
        ):

            for item in st.session_state[
                "round2_reasons"
            ]:
                st.write(
                    f"• {item}"
                )

        st.markdown(
            "### 🧠 Phân tích sơ bộ"
        )

        for finding in st.session_state[
            "round2_findings"
        ]:
            st.write(
                f"• {finding}"
            )

        st.markdown(
            "### 💡 Khuyến nghị cho người được sàng lọc"
        )

        for item in st.session_state[
            "round2_advice"
        ]:
            st.write(
                f"→ {item}"
            )

        st.caption(
            "Những nhận định trên chỉ hỗ trợ sàng lọc và phải được "
            "đối chiếu với lâm sàng, xét nghiệm chuyên sâu và nhân viên y tế."
        )

        # ----------------------------------------------------
        # MEDICAL FACILITIES
        # ----------------------------------------------------

        st.subheader(
            "🏥 CƠ SỞ Y TẾ GỢI Ý"
        )

        facilities = recommended_facilities(patient["province"])

        if facilities:

            st.success(
                f"Đã tìm thấy {len(facilities)} cơ sở ưu tiên trong "
                f"**{patient['province']}**."
            )

            st.caption(
                "Ưu tiên bệnh viện hạng I và bệnh viện tuyến Trung ương/hạng đặc biệt "
                "đang có trong danh mục prototype của tỉnh/thành. "
                "Khi đi khám, người bệnh nên hỏi trước khoa Huyết học/Truyền máu "
                "và khả năng thực hiện xét nghiệm chuyên sâu."
            )

            for i, facility in enumerate(facilities, start=1):
                with st.container(border=True):
                    st.markdown(
                        f"### {i}. {facility['name']}"
                    )
                    st.write(
                        f"**Phân loại:** {facility['tier']}"
                    )
                    st.write(
                        f"**Gợi ý:** {facility['note']}"
                    )
                    st.link_button(
                        "🗺️ Xem vị trí / chỉ đường",
                        facility["maps"],
                    )

        else:

            st.info(
                f"Prototype chưa có danh mục bệnh viện ưu tiên cho **{patient['province']}**. "
                "Bạn có thể dùng Google Maps để tìm bệnh viện hạng I hoặc cơ sở tuyến Trung ương "
                "trong chính tỉnh/thành."
            )

            maps_query = (
                f"bệnh viện hạng I bệnh viện Trung ương {patient['province']}"
            )
            maps_url = (
                "https://www.google.com/maps/search/?api=1&query="
                + requests.utils.quote(maps_query)
            )
            st.link_button(
                "🗺️ Tìm bệnh viện tuyến trên trong tỉnh",
                maps_url,
            )

        st.caption(
            "Danh mục cơ sở dùng cho điều hướng prototype; cần cập nhật/đối soát định kỳ "
            "với nguồn chính thức trước khi dùng trong nghiên cứu hoặc triển khai thực tế."
        )

        # ----------------------------------------------------
        # WORD
        # ----------------------------------------------------

        st.subheader(
            "📄 PHIẾU KẾT QUẢ"
        )

        report = make_word(
            patient=patient,
            r1_score=st.session_state[
                "round1_score"
            ],
            r1_category=st.session_state[
                "round1_category"
            ],
            r1_reasons=st.session_state[
                "round1_reasons"
            ],
            r2={
                "altitude_choice": st.session_state[
                    "round2_altitude_choice"
                ],
                "adjustment": st.session_state[
                    "round2_adjustment"
                ],
                "hb": st.session_state[
                    "round2_hb"
                ],
                "hb_adjusted": st.session_state[
                    "round2_hb_adjusted"
                ],
                "mcv": st.session_state[
                    "round2_mcv_result"
                ],
                "mch": st.session_state[
                    "round2_mch_result"
                ],
                "rbc": st.session_state[
                    "round2_rbc_result"
                ],
                "rdw": st.session_state[
                    "round2_rdw_result"
                ],
                "mentzer": st.session_state[
                    "round2_mentzer"
                ],
                "score": st.session_state[
                    "round2_score"
                ],
                "category": st.session_state[
                    "round2_category"
                ],
                "conclusion": st.session_state[
                    "round2_conclusion"
                ],
                "findings": st.session_state[
                    "round2_findings"
                ],
                "advice": st.session_state[
                    "round2_advice"
                ],
            },
        )

        st.download_button(
            "📥 TẢI PHIẾU WORD",
            data=report,
            file_name=(
                "Phieu_Thalassemia_"
                f"{safe_filename(patient['full_name'])}.docx"
            ),
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            use_container_width=True,
        )


# ============================================================
# FOOTER
# ============================================================

st.divider()

st.caption(
    "V5 Prototype — các trọng số/ngưỡng cần validation trên dữ liệu "
    "người Việt Nam trước khi sử dụng trong nghiên cứu lâm sàng."
)
